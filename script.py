import os
import docx
import boto3
from time import sleep

def readText(file):
	doc = docx.Document(file)
	fullText = ''
	for paragraph in doc.paragraphs:
		fullText += paragraph.text
		fullText += '\n\n'
	return fullText

def cutText(text):
	segments = ['']
	index = 0
	charNumber = 0 
	for char in text:
		if charNumber < 2500:
			segments[index] += char
		else:
			segments.append(char)
			index += 1
			charNumber = 1
		charNumber += 1
	return segments


def translateText(text):

	translate = boto3.client(service_name = 'translate', region_name = 'eu-west-1',
		use_ssl = True)

	result = translate.translate_text(Text = text, SourceLanguageCode = 'ru',
		TargetLanguageCode = 'en')

	return result.get('TranslatedText')

def createDocx(text):
	doc = docx.Document()
	doc.add_paragraph(text)
	doc.save('file.docx')


def main():
	text = readText('Input/file.docx')
	segments = cutText(text)
	translatedText = ''
	for number, segment in enumerate(segments):
		translatedText += translateText(segment)
		sleep(20)
		print(str(number + 1) + '/' + str(len(segments)))
	createDocx(translatedText)
	print("Done")

main()
#createDocx(translateText(readtxt('Input/file.docx')))

